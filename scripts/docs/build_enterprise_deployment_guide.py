#!/usr/bin/env python3
"""Build the enterprise deployment DOCX from the canonical Markdown runbook.

The Markdown runbook remains authoritative. This exporter intentionally supports
the Markdown constructs used by that runbook and applies a deterministic Word
style system suitable for an operator reference guide.
"""

from __future__ import annotations

import argparse
import hashlib
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_SOURCE = REPO_ROOT / "playbooks" / "enterprise-client-deployment.md"
DEFAULT_OUTPUT = (
    REPO_ROOT / "docs" / "deployment" / "Scanalyze_Enterprise_Deployment_Guide.docx"
)

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
NAVY = "203748"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
RISK_RED = "9B1C1C"
RISK_FILL = "FDECEC"
GOLD = "8A6500"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    return parser.parse_args()


def clean_text(value: str) -> str:
    return (
        value.replace("\u00a0", " ")
        .replace("\u2011", "-")
        .replace("\u2013", "-")
        .replace("\u2014", " - ")
        .replace("\u2018", "'")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u2192", "->")
    )


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def apply_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must total {CONTENT_WIDTH_DXA}: {widths}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "C9D2DC")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def choose_table_widths(rows: list[list[str]]) -> list[int]:
    columns = max(len(row) for row in rows)
    lengths = []
    for index in range(columns):
        longest = max(len(row[index]) if index < len(row) else 0 for row in rows)
        lengths.append(max(8, min(longest, 48)))
    total = sum(lengths)
    widths = [max(800, round(CONTENT_WIDTH_DXA * length / total)) for length in lengths]
    difference = CONTENT_WIDTH_DXA - sum(widths)
    widths[-1] += difference
    if widths[-1] < 800:
        shortage = 800 - widths[-1]
        widths[-1] = 800
        largest = max(range(len(widths) - 1), key=lambda i: widths[i])
        widths[largest] -= shortage
    return widths


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in doc.styles:
        style = doc.styles.add_style("Code Block", 1)
    else:
        style = doc.styles["Code Block"]
    style.font.name = "Consolas"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    style.font.size = Pt(8.5)
    style.paragraph_format.space_before = Pt(4)
    style.paragraph_format.space_after = Pt(7)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.left_indent = Inches(0.12)
    style.paragraph_format.right_indent = Inches(0.08)


def add_numbering_definition(doc: Document, ordered: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    bullet_markers = ("•", "◦", "▪")
    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), f"%{level + 1}." if ordered else bullet_markers[level])
        lvl.append(lvl_text)
        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "tab")
        lvl.append(suffix)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        text_indent = 540 + (level * 360)
        tab.set(qn("w:pos"), str(text_indent))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(text_indent))
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
    numbering.append(abstract)

    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(level, 2)))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend((color, underline))
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = clean_text(text)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(
    r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\)|(?<!\*)\*[^*]+\*(?!\*))"
)


def add_inline(paragraph, text: str, *, size: float | None = None) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(clean_text(text[cursor : match.start()]))
            set_run_font(run, size=size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(clean_text(token[2:-2]))
            set_run_font(run, size=size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(clean_text(token[1:-1]))
            set_run_font(run, name="Consolas", size=(size or 11) - 0.5, color=DARK_BLUE)
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        else:
            run = paragraph.add_run(clean_text(token[1:-1]))
            set_run_font(run, size=size, italic=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(clean_text(text[cursor:]))
        set_run_font(run, size=size)


def add_paragraph_fill(paragraph, fill: str, border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border:
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "14")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        borders.append(left)
        p_pr.append(borders)


def add_page_field(paragraph) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = OxmlElement("w:r")
    run.extend((begin, instruction, separate, value, end))
    paragraph._p.append(run)


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), WD_TAB_ALIGNMENT.RIGHT
    )
    left = paragraph.add_run("SCANALYZE | ENTERPRISE DEPLOYMENT GUIDE")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = paragraph.add_run("\tDRAFT / NO-GO")
    set_run_font(right, size=8.5, color=RISK_RED, bold=True)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), WD_TAB_ALIGNMENT.RIGHT
    )
    left = paragraph.add_run("Internal | Sanitized operational reference")
    set_run_font(left, size=8, color=MUTED)
    right = paragraph.add_run("\tPage ")
    set_run_font(right, size=8, color=MUTED)
    add_page_field(paragraph)


def extract_metadata(markdown: str) -> tuple[str, str, str]:
    version = re.search(r"\*\*Version:\*\*\s*(.+)", markdown)
    date = re.search(r"\*\*Date:\*\*\s*(.+)", markdown)
    status = re.search(r"\*\*Status:\*\*\s*(.+)", markdown)
    return (
        clean_text(version.group(1).strip()) if version else "4.0",
        clean_text(date.group(1).strip()) if date else "2026-07-10",
        clean_text(status.group(1).strip()) if status else "DRAFT / NON-EXECUTABLE / NO-GO",
    )


def add_cover(doc: Document, version: str, date: str, status: str, source_sha: str) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(96)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("ENTERPRISE DEPLOYMENT OPERATIONS")
    set_run_font(run, size=10, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Scanalyze Enterprise\nDeployment Guide")
    set_run_font(run, size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    run = subtitle.add_run(
        "Account-per-Deployment | Monorepo | AWS | Terraform | OCI supply chain"
    )
    set_run_font(run, size=13, color=DARK_BLUE)

    status_box = doc.add_paragraph()
    status_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status_box.paragraph_format.left_indent = Inches(0.35)
    status_box.paragraph_format.right_indent = Inches(0.35)
    status_box.paragraph_format.space_before = Pt(8)
    status_box.paragraph_format.space_after = Pt(24)
    add_paragraph_fill(status_box, RISK_FILL, RISK_RED)
    run = status_box.add_run(status)
    set_run_font(run, size=11, color=RISK_RED, bold=True)

    metadata = (
        ("Version", version),
        ("Updated", date),
        ("Classification", "Internal - sanitized operational reference"),
        ("Canonical source", "playbooks/enterprise-client-deployment.md"),
        ("Source SHA-256", source_sha),
    )
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(f"{label}: ")
        set_run_font(run, size=9.5, color=MUTED, bold=True)
        run = paragraph.add_run(value)
        set_run_font(run, name="Consolas" if label == "Source SHA-256" else "Calibri", size=9.5, color=MUTED)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(24)
    note.paragraph_format.left_indent = Inches(0.55)
    note.paragraph_format.right_indent = Inches(0.55)
    add_inline(
        note,
        "Evidence before claims: this document does not authorize AWS mutation. "
        "Every blocked capability must be implemented, validated in non-production "
        "and approved before production execution.",
        size=9.5,
    )
    note.runs[0].italic = True
    doc.add_page_break()


def markdown_table_rows(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [clean_text(cell.strip()) for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    normalized = [row + [""] * (columns - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=columns)
    widths = choose_table_widths(normalized)
    for row_index, row in enumerate(normalized):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.line_spacing = 1.05
            add_inline(paragraph, value, size=8.5 if columns >= 4 else 9)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
                shade_cell(cell, LIGHT_BLUE)
    set_repeat_table_header(table.rows[0])
    apply_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph(style="Code Block")
    add_paragraph_fill(paragraph, LIGHT_GRAY)
    text = "\n".join(clean_text(line) for line in lines).rstrip()
    run = paragraph.add_run(text)
    set_run_font(run, name="Consolas", size=8.5, color=NAVY)


def add_static_contents(doc: Document, markdown_lines: list[str]) -> None:
    doc.add_heading("Contents", level=1)
    paragraph = doc.add_paragraph()
    add_inline(
        paragraph,
        "This navigation list mirrors the canonical Markdown headings. Word's "
        "Navigation Pane also uses the embedded heading hierarchy.",
    )
    for line in markdown_lines:
        match = re.match(r"^##\s+(.+)$", line)
        if not match or match.group(1).strip().lower() in {
            "table of contents",
            "tabla de contenido",
            "contents",
        }:
            continue
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.space_after = Pt(4)
        add_inline(paragraph, match.group(1).strip())
    doc.add_page_break()


def add_markdown_body(doc: Document, markdown: str, bullet_num_id: int, ordered_num_id: int) -> None:
    lines = markdown.splitlines()
    index = 0
    skip_contents = False
    paragraph_buffer: list[str] = []
    active_ordered_num_id: int | None = None

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        paragraph = doc.add_paragraph()
        add_inline(paragraph, " ".join(part.strip() for part in paragraph_buffer))
        paragraph_buffer.clear()

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.lower() in {
            "## table of contents",
            "## tabla de contenido",
            "## contents",
        }:
            flush_paragraph()
            active_ordered_num_id = None
            skip_contents = True
            index += 1
            continue
        if skip_contents:
            if stripped.startswith("## "):
                skip_contents = False
            else:
                index += 1
                continue

        if not stripped:
            flush_paragraph()
            active_ordered_num_id = None
            index += 1
            continue
        if stripped in {"---", "***"}:
            flush_paragraph()
            active_ordered_num_id = None
            index += 1
            continue
        if stripped.startswith("# "):
            flush_paragraph()
            active_ordered_num_id = None
            index += 1
            continue
        if re.match(r"^\*\*(Version|Date|Status|Audience|Owner|Classification|Canonical source):", stripped):
            flush_paragraph()
            active_ordered_num_id = None
            index += 1
            continue
        if stripped.startswith("```"):
            flush_paragraph()
            active_ordered_num_id = None
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(doc, code_lines)
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and lines[index + 1].strip().startswith("|"):
            flush_paragraph()
            active_ordered_num_id = None
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(doc, markdown_table_rows(table_lines))
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            active_ordered_num_id = None
            level = len(heading.group(1)) - 1
            title = clean_text(heading.group(2).strip())
            if title.lower().startswith("appendix"):
                doc.add_page_break()
            paragraph = doc.add_heading(level=min(level, 3))
            add_inline(paragraph, title)
            index += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            active_ordered_num_id = None
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip().lstrip(">").strip())
                index += 1
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.right_indent = Inches(0.12)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            add_paragraph_fill(paragraph, CALLOUT, BLUE)
            if quote_lines:
                quote_lines[0] = re.sub(
                    r"^\[!(CAUTION|IMPORTANT|WARNING|NOTE)\]\s*",
                    lambda match: f"**{match.group(1).title()}.** ",
                    quote_lines[0],
                )
            add_inline(paragraph, " ".join(quote_lines))
            continue

        checklist = re.match(r"^(\s*)[-*+]\s+\[([ xX])\]\s+(.+)$", line)
        if checklist:
            flush_paragraph()
            active_ordered_num_id = None
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25 + len(checklist.group(1)) / 12)
            paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            paragraph.paragraph_format.space_after = Pt(4)
            marker = "[x] " if checklist.group(2).lower() == "x" else "[ ] "
            run = paragraph.add_run(marker)
            set_run_font(run, name="Consolas", size=10, color=DARK_BLUE, bold=True)
            add_inline(paragraph, checklist.group(3))
            index += 1
            continue

        bullet = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        if bullet:
            flush_paragraph()
            active_ordered_num_id = None
            item_text = bullet.group(2)
            index += 1
            while index < len(lines):
                continuation = lines[index]
                continuation_stripped = continuation.strip()
                if (
                    not continuation_stripped
                    or re.match(r"^(#{1,4})\s+", continuation_stripped)
                    or continuation_stripped.startswith(("```", ">", "|"))
                    or re.match(r"^\s*[-*+]\s+", continuation)
                    or re.match(r"^\s*\d+[.)]\s+", continuation)
                ):
                    break
                item_text += " " + continuation_stripped
                index += 1
            paragraph = doc.add_paragraph()
            apply_numbering(paragraph, bullet_num_id, len(bullet.group(1)) // 2)
            add_inline(paragraph, item_text)
            continue

        ordered = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
        if ordered:
            flush_paragraph()
            if active_ordered_num_id is None:
                active_ordered_num_id = add_numbering_definition(doc, ordered=True)
            item_text = ordered.group(2)
            index += 1
            while index < len(lines):
                continuation = lines[index]
                continuation_stripped = continuation.strip()
                if (
                    not continuation_stripped
                    or re.match(r"^(#{1,4})\s+", continuation_stripped)
                    or continuation_stripped.startswith(("```", ">", "|"))
                    or re.match(r"^\s*[-*+]\s+", continuation)
                    or re.match(r"^\s*\d+[.)]\s+", continuation)
                ):
                    break
                item_text += " " + continuation_stripped
                index += 1
            paragraph = doc.add_paragraph()
            apply_numbering(paragraph, active_ordered_num_id, len(ordered.group(1)) // 2)
            add_inline(paragraph, item_text)
            continue

        active_ordered_num_id = None
        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()


def enable_field_updates(doc: Document) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def build(source: Path, output: Path) -> None:
    markdown_bytes = source.read_bytes()
    markdown = markdown_bytes.decode("utf-8")
    source_sha = hashlib.sha256(markdown_bytes).hexdigest()
    version, date, status = extract_metadata(markdown)

    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
    enable_field_updates(doc)

    properties = doc.core_properties
    properties.title = "Scanalyze Enterprise Deployment Guide"
    properties.subject = "Enterprise deployment governance and operational runbook"
    properties.author = "Scanalyze Platform Engineering"
    properties.last_modified_by = "Scanalyze Platform Engineering"
    properties.keywords = "Scanalyze, AWS, Terraform, monorepo, deployment, security"

    bullet_num_id = add_numbering_definition(doc, ordered=False)
    ordered_num_id = add_numbering_definition(doc, ordered=True)

    add_cover(doc, version, date, status, source_sha)
    add_static_contents(doc, markdown.splitlines())
    add_markdown_body(doc, markdown, bullet_num_id, ordered_num_id)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Wrote {output}")
    print(f"Source SHA-256: {source_sha}")


def main() -> None:
    args = parse_args()
    build(args.source.resolve(), args.output.resolve())


if __name__ == "__main__":
    main()
